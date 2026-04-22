import tkinter as tk
from tkinter import messagebox, ttk
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import os

# --- HÀM CHUYỂN SỐ THÀNH CHỮ TIẾNG VIỆT ---
def number_to_vietnamese(number):
    units = ["", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]
    
    def read_three_digits(n, show_zero_hundred=False):
        res = ""
        h = n // 100
        t = (n % 100) // 10
        u = n % 10
        if h > 0 or show_zero_hundred:
            res += units[h] + " trăm "
        if t == 0:
            if h > 0 and u > 0: res += "lẻ "
        elif t == 1:
            res += "mười "
        else:
            res += units[t] + " mươi "
        if t > 1 and u == 1:
            res += "mốt"
        elif t > 0 and u == 5:
            res += "lăm"
        elif u > 0:
            res += units[u]
        return res.strip()

    if number == 0: return "Không đồng"
    res = ""
    triệu = number // 1000000
    nghìn = (number % 1000000) // 1000
    đồng = number % 1000
    
    if triệu > 0:
        res += read_three_digits(triệu) + " triệu "
    if nghìn > 0:
        res += read_three_digits(nghìn, triệu > 0) + " nghìn "
    if đồng > 0:
        res += read_three_digits(đồng, nghìn > 0) + " đồng"
    else:
        res += " đồng"
    
    res = res.strip().capitalize() + " chẵn."
    return res.replace("  ", " ")

class InvoiceApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Phần mềm Hóa Đơn Tuyên Ninh - HUST SEEE")
        self.root.geometry("900x750")
        self.items = []

        # --- GIAO DIỆN NHẬP LIỆU ---
        info_frame = tk.LabelFrame(root, text="Thông tin hóa đơn", padx=10, pady=10)
        info_frame.pack(fill="x", padx=15, pady=5)

        tk.Label(info_frame, text="Tên Cửa hàng:").grid(row=0, column=0, sticky="w")
        self.ent_store = tk.Entry(info_frame, width=30)
        self.ent_store.insert(0, "TUYÊN NINH")
        self.ent_store.grid(row=0, column=1, pady=2)

        tk.Label(info_frame, text="Khách hàng:").grid(row=0, column=2, sticky="w", padx=20)
        self.ent_cust = tk.Entry(info_frame, width=30)
        self.ent_cust.grid(row=0, column=3, pady=2)

        item_frame = tk.LabelFrame(root, text="Chi tiết mặt hàng", padx=10, pady=10)
        item_frame.pack(fill="x", padx=15, pady=5)

        fields = ["Tên hàng", "ĐVT", "SL", "Đơn giá"]
        self.entries = {}
        for i, f in enumerate(fields):
            tk.Label(item_frame, text=f).grid(row=0, column=i)
            e = tk.Entry(item_frame, width=18)
            e.grid(row=1, column=i, padx=5)
            self.entries[f] = e
            # BỔ SUNG: Nhấn Enter ở bất kỳ ô nào cũng sẽ thêm dòng
            e.bind('<Return>', lambda event: self.add_item())

        btn_add = tk.Button(item_frame, text="Thêm dòng (Enter)", command=self.add_item, bg="#4CAF50", fg="white")
        btn_add.grid(row=1, column=4, padx=10)

        # --- BẢNG HIỂN THỊ ---
        self.tree = ttk.Treeview(root, columns=("1", "2", "3", "4", "5"), show='headings', height=12)
        for id, txt, w in [("1", "Tên hàng", 300), ("2", "ĐVT", 80), ("3", "SL", 60), ("4", "Đơn giá", 120), ("5", "Thành tiền", 150)]:
            self.tree.heading(id, text=txt)
            self.tree.column(id, width=w)
        self.tree.pack(fill="both", expand=True, padx=15, pady=10)

        # --- TỔNG TIỀN TỰ ĐỘNG ---
        bottom_frame = tk.Frame(root)
        bottom_frame.pack(fill="x", padx=15, pady=10)
        self.lbl_total = tk.Label(bottom_frame, text="TỔNG: 0 VNĐ", font=("Arial", 14, "bold"), fg="red")
        self.lbl_total.pack(side="right")

        tk.Label(bottom_frame, text="Bằng chữ:").pack(side="left")
        self.ent_words = tk.Entry(bottom_frame, width=65, font=("Arial", 10, "italic"))
        self.ent_words.pack(side="left", padx=10)

        tk.Button(root, text="XUẤT HÓA ĐƠN & MỞ FILE", command=self.export_word, bg="#d32f2f", fg="white", font=("Arial", 12, "bold"), height=2).pack(fill="x", padx=150, pady=15)

    def add_item(self):
        try:
            name = self.entries["Tên hàng"].get()
            unit = self.entries["ĐVT"].get()
            qty = float(self.entries["SL"].get())
            price = float(self.entries["Đơn giá"].get())
            total = qty * price
            
            if not name: return
            
            self.items.append([name, unit, qty, price, total])
            self.tree.insert("", "end", values=(name, unit, qty, f"{price:,.0f}", f"{total:,.0f}"))
            
            # Cập nhật tổng số tiền
            grand_total = int(sum(item[4] for item in self.items))
            self.lbl_total.config(text=f"TỔNG: {grand_total:,.0f} VNĐ")
            
            # TỰ ĐỘNG ĐIỀN CHỮ
            self.ent_words.delete(0, tk.END)
            self.ent_words.insert(0, number_to_vietnamese(grand_total))
            
            # Xóa các ô nhập để nhập tiếp, focus lại vào ô Tên hàng
            for e in self.entries.values(): e.delete(0, tk.END)
            self.entries["Tên hàng"].focus_set()
            
        except ValueError:
            messagebox.showerror("Lỗi", "Vui lòng nhập SL và Đơn giá bằng số!")

    def export_word(self):
        if not self.items: return
        try:
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            red_color = RGBColor(192, 0, 0)

            # Cửa hàng
            p = doc.add_paragraph()
            r = p.add_run(self.ent_store.get().upper())
            r.bold = True
            r.font.color.rgb = red_color

            # Tiêu đề
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rt = title.add_run("\nHÓA ĐƠN BÁN HÀNG")
            rt.bold = True
            rt.font.size = Pt(20)
            rt.font.color.rgb = red_color

            doc.add_paragraph(f"Khách hàng: {self.ent_cust.get()}")
            
            # Bảng giống mẫu
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, t in enumerate(["STT", "Tên hàng", "ĐVT", "SL", "Đơn giá", "Thành tiền"]):
                hdr[i].text = t
                hdr[i].paragraphs[0].runs[0].font.bold = True

            total_val = 0
            for idx, item in enumerate(self.items, 1):
                row = table.add_row().cells
                row[0].text = str(idx)
                row[1].text = item[0]
                row[2].text = item[1]
                row[3].text = str(int(item[2]) if item[2].is_integer() else item[2])
                row[4].text = f"{item[3]:,.0f}"
                row[5].text = f"{item[4]:,.0f}"
                total_val += item[4]

            row_sum = table.add_row().cells
            row_sum[0].merge(row_sum[4])
            row_sum[0].text = "TỔNG CỘNG"
            row_sum[5].text = f"{total_val:,.0f}"
            row_sum[5].paragraphs[0].runs[0].font.bold = True

            doc.add_paragraph(f"\nSố tiền bằng chữ: {self.ent_words.get()}")
            doc.add_paragraph(f"Ngày {datetime.datetime.now().day} tháng {datetime.datetime.now().month} năm {datetime.datetime.now().year}").alignment = WD_ALIGN_PARAGRAPH.RIGHT

            fname = f"HoaDon_{datetime.datetime.now().strftime('%H%M%S')}.docx"
            doc.save(fname)
            if messagebox.askyesno("Thành công", f"Đã xuất hóa đơn!\nBạn có muốn mở file ngay không?"):
                os.startfile(fname)
        except Exception as e:
            messagebox.showerror("Lỗi", f"Lỗi: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = InvoiceApp(root)
    root.mainloop()